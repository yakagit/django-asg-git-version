from docx import Document
import os


def GenerateDocx(data, template_path='X:/Python/call_template.docx'):
    doc = Document(template_path)

    replacements = {
        "{{nadzor_num}}": data['nadzor_num'],
        "{{nadzor_date}}": data['nadzor_date'],
        "{{works}}": data['works'],
        "{{interval}}": data['interval'],
        "{{diameter}}": data['diameter'],
        "{{length}}": data['length'],
        "{{material}}": data['material'],
        "{{agent}}": data['agent'],
        "{{agent_phone}}": data['agent_phone'],
        "{{meet_date}}": data['meet_date'],
        "{{meet_place}}": data['meet_place'],
        "{{zone}}": data['zone'],
        "{{object_full_name}}": data['object_full_name'],
    }

    # Проходим по каждому параграфу и заменяем ключевые слова на значения из словаря replacements
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, str(value))

    # Проходим по каждому заголовку таблицы и заменяем ключевые слова в таблицах
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, str(value))

    # Сохраняем временный файл
    disk_path = 'X:/Python/'
    output_path = disk_path + f'generated_{data["object_short_name"]}.docx'
    doc.save(output_path)
    
    return output_path
